import json
import logging
import re
from io import BytesIO

from docx import Document
from fastapi import HTTPException, UploadFile
from openai import AsyncOpenAI
from pydantic import ValidationError
from pypdf import PdfReader

from backend.app.core.config import settings
from backend.app.prompts.resume.resume_import_prompts import (
    RESUME_IMPORT_SYSTEM_PROMPT,
    build_resume_import_prompt,
)
from backend.app.schemas.resume.resume_import import ResumeImportResponse


client = AsyncOpenAI(
    api_key=settings.OPENAI_API_KEY,
)

logger = logging.getLogger(__name__)


MAX_FILE_SIZE = 5 * 1024 * 1024

MAX_TEXT_LENGTH = 15000

MIN_TEXT_LENGTH = 50

PDF_CONTENT_TYPES = {"application/pdf"}

DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

WHITESPACE_RE = re.compile(r"[ \t]+")
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _detect_file_kind(file: UploadFile) -> str:
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    if content_type in PDF_CONTENT_TYPES or filename.endswith(".pdf"):
        return "pdf"

    if content_type in DOCX_CONTENT_TYPES or filename.endswith(".docx"):
        return "docx"

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Please upload a PDF or DOCX file.",
    )


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))

        pages_text = [
            page.extract_text() or ""
            for page in reader.pages
        ]
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="Could not read this PDF. Please upload a valid, non-encrypted PDF file.",
        ) from exc

    return "\n".join(pages_text)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
        ]

        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="Could not read this DOCX file. Please upload a valid Word document.",
        ) from exc

    return "\n".join(paragraphs + table_cells)


def _clean_extracted_text(text: str) -> str:
    text = CONTROL_CHARS_RE.sub("", text)
    text = WHITESPACE_RE.sub(" ", text)

    lines = [
        line.strip()
        for line in text.splitlines()
    ]

    cleaned = "\n".join(line for line in lines if line)

    return cleaned[:MAX_TEXT_LENGTH]


async def import_resume_from_file(
    *,
    file: UploadFile,
    user_id: int,
) -> ResumeImportResponse:
    """
    Extract structured resume data from an uploaded PDF/DOCX file using an
    LLM, so the Resume Builder form can be auto-filled.

    Note: unlike `resume_analysis_service.py`, this intentionally does NOT
    run the extracted text through `llm_privacy.prepare_data()` — that
    sanitizer strips exactly the contact fields (name, email, phone, links)
    this feature is meant to extract. The user is uploading their own resume
    for the explicit purpose of populating those fields.
    """
    file_kind = _detect_file_kind(file)

    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(
            status_code=400,
            detail="The uploaded file is empty.",
        )

    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail="File is too large. Maximum size is 5MB.",
        )

    if file_kind == "pdf":
        raw_text = _extract_text_from_pdf(file_bytes)
    else:
        raw_text = _extract_text_from_docx(file_bytes)

    clean_text = _clean_extracted_text(raw_text)

    if len(clean_text) < MIN_TEXT_LENGTH:
        raise HTTPException(
            status_code=400,
            detail=(
                "Could not read enough text from this file. It may be a "
                "scanned image, which isn't supported yet."
            ),
        )

    prompt = build_resume_import_prompt(raw_text=clean_text)

    try:
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.2,
            response_format={
                "type": "json_object",
            },
            messages=[
                {
                    "role": "system",
                    "content": RESUME_IMPORT_SYSTEM_PROMPT,
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )
    except Exception as exc:
        logger.exception(
            "Resume import AI request failed for user_id=%s file_kind=%s text_length=%s",
            user_id,
            file_kind,
            len(clean_text),
        )
        raise HTTPException(
            status_code=500,
            detail="Failed to extract data from resume",
        ) from exc

    content = response.choices[0].message.content or "{}"

    try:
        parsed = json.loads(content)
    except Exception as exc:
        logger.exception(
            "Resume import JSON parsing failed for user_id=%s file_kind=%s",
            user_id,
            file_kind,
        )
        raise HTTPException(
            status_code=500,
            detail="Invalid response format",
        ) from exc

    try:
        result = ResumeImportResponse(**parsed)
    except ValidationError as exc:
        logger.exception(
            "Resume import response validation failed for user_id=%s file_kind=%s",
            user_id,
            file_kind,
        )
        raise HTTPException(
            status_code=500,
            detail="Unexpected AI response format",
        ) from exc

    return _drop_empty_placeholders(result)


def _non_empty(values: list[str]) -> list[str]:
    return [value for value in values if value.strip()]


def _drop_empty_placeholders(
    result: ResumeImportResponse,
) -> ResumeImportResponse:
    """
    Some LLM responses include placeholder entries (e.g. a project with all
    empty fields, or a bullet list containing a single empty string) instead
    of omitting the section entirely. Strip those out so the frontend never
    renders blank cards/bullets after an import.
    """
    result.experience = [
        item
        for item in result.experience
        if item.company.strip() or item.role.strip()
    ]

    for item in result.experience:
        item.bullets = _non_empty(item.bullets)

    result.education = [
        item
        for item in result.education
        if item.school.strip() or item.degree.strip()
    ]

    for item in result.education:
        item.bullets = _non_empty(item.bullets)

    result.projects = [
        item
        for item in result.projects
        if item.title.strip() or item.description.strip()
    ]

    for item in result.projects:
        item.bullets = _non_empty(item.bullets)
        item.technologies = _non_empty(item.technologies)

    result.skills = [
        item
        for item in result.skills
        if _non_empty(item.skills)
    ]

    for item in result.skills:
        item.skills = _non_empty(item.skills)

    result.languages = [
        item
        for item in result.languages
        if item.name.strip()
    ]

    result.softSkills = _non_empty(result.softSkills)

    return result
